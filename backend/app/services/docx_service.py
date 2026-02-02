from __future__ import annotations

from datetime import datetime
from pathlib import Path

from docx import Document
from docx.shared import Pt

from ..utils.filesystem import build_tree, gather_file_records


def add_code_block(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def generate_docx(root_dir: Path, output_path: Path) -> int:
    document = Document()
    document.add_heading("Project Documentation Export", level=1)
    document.add_paragraph(f"Source path: {root_dir}")
    document.add_paragraph(f"Generated at: {datetime.utcnow().isoformat()} UTC")

    document.add_heading("Folder Structure", level=2)
    tree_text = build_tree(root_dir)
    tree_paragraph = document.add_paragraph()
    add_code_block(tree_paragraph, tree_text)

    document.add_page_break()
    document.add_heading("File Contents", level=2)

    file_records = gather_file_records(root_dir)
    for record in file_records:
        document.add_heading(record["path"], level=3)
        document.add_paragraph(
            f"Size: {record['size']} bytes | Modified: {record['modified']}"
        )
        content_paragraph = document.add_paragraph()
        add_code_block(content_paragraph, record["content"])

    document.save(output_path)
    return len(file_records)
