from __future__ import annotations

import io
import json
import os
import shutil
import tempfile
from concurrent.futures import ThreadPoolExecutor, as_completed
from datetime import datetime
from pathlib import Path
from typing import Iterable

from docx import Document
from docx.shared import Pt
from fastapi import FastAPI, File, Form, HTTPException, UploadFile
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import FileResponse, JSONResponse
from fastapi.staticfiles import StaticFiles
from pydantic import BaseModel, Field

PROJECT_ROOT = Path(__file__).resolve().parent.parent
FRONTEND_DIR = PROJECT_ROOT / "frontend"
STORAGE_DIR = PROJECT_ROOT / "shared" / "runs"
STORAGE_DIR.mkdir(parents=True, exist_ok=True)

MAX_WORKERS = max(os.cpu_count() or 2, 4)

app = FastAPI(title="Project Analyzer Bot", version="1.0.0")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"]
)


class PromptRequest(BaseModel):
    task: str = Field(..., min_length=5, description="User's task description")
    project_summary: str | None = Field(
        None,
        description="Optional summary of the project to include in the prompt"
    )


class ApplyOperation(BaseModel):
    action: str = Field(..., description="create|update|delete")
    path: str = Field(..., description="Relative path inside the project")
    content: str | None = Field(None, description="New file contents for create/update")


class ApplyRequest(BaseModel):
    project_root: str = Field(..., description="Absolute path to project root")
    operations: list[ApplyOperation]


class AnalyzeResponse(BaseModel):
    run_id: str
    docx_path: str
    file_count: int


EXCLUDED_DIRS = {".git", ".idea", ".vscode", "node_modules", "target", "build"}


def is_binary_file(path: Path) -> bool:
    try:
        with path.open("rb") as handle:
            chunk = handle.read(2048)
        return b"\x00" in chunk
    except OSError:
        return True


def iter_files(root_dir: Path) -> Iterable[Path]:
    for base, dirs, files in os.walk(root_dir):
        dirs[:] = [d for d in dirs if d not in EXCLUDED_DIRS]
        for filename in files:
            yield Path(base) / filename


def build_tree(root_dir: Path) -> str:
    lines: list[str] = []
    root_dir = root_dir.resolve()
    prefix_map: dict[Path, str] = {root_dir: ""}

    for base, dirs, files in os.walk(root_dir):
        base_path = Path(base)
        prefix = prefix_map.get(base_path, "")
        entries = sorted(dirs) + sorted(files)
        for index, name in enumerate(entries):
            path = base_path / name
            connector = "└── " if index == len(entries) - 1 else "├── "
            lines.append(f"{prefix}{connector}{name}")
            if path.is_dir():
                extension = "    " if index == len(entries) - 1 else "│   "
                prefix_map[path] = prefix + extension
    return "\n".join(lines)


def read_file_text(path: Path) -> str:
    if is_binary_file(path):
        return "<binary file omitted>"
    try:
        return path.read_text(encoding="utf-8", errors="replace")
    except OSError as exc:
        return f"<unable to read file: {exc}>"


def gather_file_records(root_dir: Path) -> list[dict[str, str]]:
    file_paths = list(iter_files(root_dir))
    records: list[dict[str, str]] = []
    with ThreadPoolExecutor(max_workers=MAX_WORKERS) as executor:
        futures = {executor.submit(read_file_text, path): path for path in file_paths}
        for future in as_completed(futures):
            path = futures[future]
            relative_path = path.relative_to(root_dir).as_posix()
            stat = path.stat()
            content = future.result()
            records.append(
                {
                    "path": relative_path,
                    "size": str(stat.st_size),
                    "modified": datetime.fromtimestamp(stat.st_mtime).isoformat(),
                    "content": content,
                }
            )
    records.sort(key=lambda item: item["path"])
    return records


def add_code_block(paragraph, text: str) -> None:
    run = paragraph.add_run(text)
    run.font.name = "Consolas"
    run.font.size = Pt(9)


def generate_docx(root_dir: Path, output_path: Path) -> int:
    document = Document()
    document.add_heading("Project Documentation Export", level=1)
    document.add_paragraph(f"Source path: {root_dir}")
    document.add_paragraph(f"Generated at: {datetime.utcnow().isoformat()} UTC")

    document.add_heading("Folder Structure", level=2)
    tree_text = build_tree(root_dir)
    tree_paragraph = document.add_paragraph()
    add_code_block(tree_paragraph, tree_text)

    document.add_page_break()
    document.add_heading("File Contents", level=2)

    file_records = gather_file_records(root_dir)
    for record in file_records:
        document.add_heading(record["path"], level=3)
        document.add_paragraph(
            f"Size: {record['size']} bytes | Modified: {record['modified']}"
        )
        content_paragraph = document.add_paragraph()
        add_code_block(content_paragraph, record["content"])

    document.save(output_path)
    return len(file_records)


def safe_join(root_dir: Path, relative_path: str) -> Path:
    destination = (root_dir / relative_path).resolve()
    if root_dir not in destination.parents and destination != root_dir:
        raise ValueError("Invalid path outside project root")
    return destination


def apply_operations(root_dir: Path, operations: list[ApplyOperation]) -> list[str]:
    messages: list[str] = []
    for operation in operations:
        action = operation.action.lower()
        destination = safe_join(root_dir, operation.path)
        if action in {"create", "update"}:
            destination.parent.mkdir(parents=True, exist_ok=True)
            if operation.content is None:
                raise ValueError(f"Missing content for {action} {operation.path}")
            destination.write_text(operation.content, encoding="utf-8")
            messages.append(f"{action}: {operation.path}")
        elif action == "delete":
            if destination.exists():
                if destination.is_dir():
                    shutil.rmtree(destination)
                else:
                    destination.unlink()
                messages.append(f"deleted: {operation.path}")
            else:
                messages.append(f"skipped (missing): {operation.path}")
        else:
            raise ValueError(f"Unknown action: {operation.action}")
    return messages


@app.get("/api/health")
async def health() -> dict[str, str]:
    return {"status": "ok"}


@app.post("/api/projects/analyze", response_model=AnalyzeResponse)
async def analyze_project(path: str = Form(...)) -> AnalyzeResponse:
    root_dir = Path(path).expanduser().resolve()
    if not root_dir.exists() or not root_dir.is_dir():
        raise HTTPException(status_code=400, detail="Project path not found")

    run_id = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    run_dir = STORAGE_DIR / run_id
    run_dir.mkdir(parents=True, exist_ok=True)
    output_path = run_dir / "project-export.docx"

    file_count = generate_docx(root_dir, output_path)
    return AnalyzeResponse(run_id=run_id, docx_path=str(output_path), file_count=file_count)


@app.post("/api/projects/upload", response_model=AnalyzeResponse)
async def analyze_upload(files: list[UploadFile] = File(...)) -> AnalyzeResponse:
    if not files:
        raise HTTPException(status_code=400, detail="No files uploaded")

    run_id = datetime.utcnow().strftime("%Y%m%d%H%M%S")
    upload_root = STORAGE_DIR / run_id / "upload"
    upload_root.mkdir(parents=True, exist_ok=True)

    for file in files:
        relative_path = Path(file.filename)
        destination = upload_root / relative_path
        destination.parent.mkdir(parents=True, exist_ok=True)
        content = await file.read()
        destination.write_bytes(content)

    output_path = STORAGE_DIR / run_id / "project-export.docx"
    file_count = generate_docx(upload_root, output_path)
    return AnalyzeResponse(run_id=run_id, docx_path=str(output_path), file_count=file_count)


@app.get("/api/projects/download")
async def download_docx(path: str) -> FileResponse:
    docx_path = Path(path)
    if not docx_path.exists():
        raise HTTPException(status_code=404, detail="DOCX not found")
    return FileResponse(docx_path, filename=docx_path.name)


@app.post("/api/prompts/build")
async def build_prompt(request: PromptRequest) -> JSONResponse:
    instructions = (
        "You are an expert software engineer. You will receive a DOCX export that "
        "contains the full project structure and source code. Based on the user's "
        "task, output a JSON array of file operations. Each operation must be an object "
        "with: action (create|update|delete), path (relative path), and content (for "
        "create/update). Return JSON only, no markdown."
    )
    prompt_payload = {
        "task": request.task,
        "project_summary": request.project_summary,
        "output_format": "JSON",
        "instructions": instructions,
        "schema": {
            "action": "create|update|delete",
            "path": "relative/file/path.ext",
            "content": "file content for create/update"
        }
    }
    return JSONResponse(prompt_payload)


@app.post("/api/patch/apply")
async def apply_patch(request: ApplyRequest) -> JSONResponse:
    root_dir = Path(request.project_root).expanduser().resolve()
    if not root_dir.exists() or not root_dir.is_dir():
        raise HTTPException(status_code=400, detail="Project root not found")

    try:
        results = apply_operations(root_dir, request.operations)
    except ValueError as exc:
        raise HTTPException(status_code=400, detail=str(exc)) from exc

    return JSONResponse({"status": "ok", "results": results})


@app.get("/")
async def index() -> FileResponse:
    return FileResponse(FRONTEND_DIR / "index.html")


app.mount("/static", StaticFiles(directory=FRONTEND_DIR), name="static")
